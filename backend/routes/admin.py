import uuid
import os
import logging
import base64
import asyncio
from io import BytesIO
from datetime import datetime, timezone
from fastapi import APIRouter, HTTPException, Depends, Response
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.units import cm
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.admin import AdminLogin, AdminTokenResponse, AdminReportData, SendReportData, RefundRequestData
from utils.auth import hash_password, verify_password, create_admin_token, get_admin_user, seed_admin_user
from services.docx_service import generate_report_docx
from config import db, ADMIN_EMAIL
import resend

logger = logging.getLogger(__name__)
router = APIRouter()


@router.post("/admin/login", response_model=AdminTokenResponse)
async def admin_login(credentials: AdminLogin):
    # Always sync admin credentials with env vars (critical for serverless cold starts)
    try:
        await seed_admin_user()
    except Exception as e:
        logger.error(f"Auto-seed admin failed: {e}")

    admin = await db.admin_users.find_one({"username": credentials.username})
    if not admin:
        default_username = os.environ.get('ADMIN_USERNAME', 'admin')
        default_password = os.environ.get('ADMIN_PASSWORD', 'admin2026')
        if credentials.username == default_username and credentials.password == default_password:
            hashed = hash_password(credentials.password)
            await db.admin_users.insert_one({
                "username": credentials.username,
                "password_hash": hashed,
                "created_at": datetime.now(timezone.utc)
            })
            token = create_admin_token()
            return AdminTokenResponse(access_token=token)
        raise HTTPException(status_code=401, detail="Invalid admin credentials")

    if not verify_password(credentials.password, admin["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid admin credentials")

    token = create_admin_token()
    return AdminTokenResponse(access_token=token)


@router.get("/admin/analyses")
async def get_all_analyses(admin: dict = Depends(get_admin_user)):
    analyses = await db.verification_results.find(
        {},
        {"_id": 0, "id": 1, "user_id": 1, "profile_name": 1, "status": 1, "created_at": 1, "form_data": 1, "ai_analysis": 1, "admin_report": 1, "photos": 1}
    ).sort("created_at", -1).to_list(100)

    user_ids = list(set(a.get("user_id") for a in analyses if a.get("user_id")))
    users_list = await db.users.find({"id": {"$in": user_ids}}, {"_id": 0, "id": 1, "email": 1, "name": 1}).to_list(len(user_ids))
    users_map = {u["id"]: u for u in users_list}

    result = []
    for analysis in analyses:
        user = users_map.get(analysis.get("user_id"), {})
        form_data = analysis.get("form_data", {})
        photos = analysis.get("photos", []) or form_data.get("photos", [])
        result.append({
            "id": analysis.get("id"),
            "user_id": analysis.get("user_id"),
            "user_email": user.get("email", "Unknown"),
            "user_name": user.get("name", "Unknown"),
            "profile_name": analysis.get("profile_name"),
            "status": analysis.get("status", "pending"),
            "created_at": analysis.get("created_at"),
            "form_data": form_data,
            "ai_analysis": analysis.get("ai_analysis"),
            "admin_report": analysis.get("admin_report"),
            "photos": photos
        })
    return result


@router.get("/admin/analyses/{analysis_id}")
async def get_admin_analysis(analysis_id: str, admin: dict = Depends(get_admin_user)):
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    user = await db.users.find_one({"id": analysis["user_id"]}, {"_id": 0, "email": 1, "name": 1})
    return {
        "id": analysis.get("id"),
        "user_id": analysis.get("user_id"),
        "user_email": user.get("email") if user else "Unknown",
        "user_name": user.get("name") if user else "Unknown",
        "profile_name": analysis.get("profile_name"),
        "status": analysis.get("status", "pending"),
        "created_at": analysis.get("created_at"),
        "form_data": analysis.get("form_data", {}),
        "ai_analysis": analysis.get("ai_analysis"),
        "admin_report": analysis.get("admin_report"),
        "overall_score": analysis.get("overall_score", 0),
        "trust_level": analysis.get("trust_level", "pending"),
        "red_flags": analysis.get("red_flags", []),
        "recommendations": analysis.get("recommendations", [])
    }


@router.patch("/admin/analyses/{analysis_id}/status")
async def update_analysis_status(analysis_id: str, new_status: str, admin: dict = Depends(get_admin_user)):
    result = await db.verification_results.update_one(
        {"id": analysis_id},
        {"$set": {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"message": "Status updated"}


@router.post("/refund-request")
async def submit_refund_request(data: RefundRequestData):
    try:
        refund_ref = f"REF-{uuid.uuid4().hex[:8].upper()}"
        refund_doc = {
            "id": refund_ref,
            "personal_info": {
                "firstName": data.firstName, "lastName": data.lastName, "username": data.username,
                "email": data.email, "phone": data.phone, "address": data.address,
                "city": data.city, "postalCode": data.postalCode, "country": data.country
            },
            "order_info": {
                "orderReference": data.orderReference, "orderDate": data.orderDate,
                "packagePurchased": data.packagePurchased, "amountPaid": data.amountPaid
            },
            "bank_info": {
                "accountHolder": data.accountHolder, "iban": data.iban,
                "bic": data.bic, "bankName": data.bankName
            },
            "reason": data.reason,
            "additionalDetails": data.additionalDetails,
            "status": "pending",
            "submittedAt": data.submittedAt,
            "language": data.language,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.refund_requests.insert_one(refund_doc)

        admin_email = os.environ.get('ADMIN_EMAIL')
        if admin_email:
            resend.api_key = os.environ.get('RESEND_API_KEY')
            isFr = data.language == 'fr'
            reason_text = {
                'insufficient_data': 'Insufficient data for analysis' if not isFr else "Donn\u00e9es insuffisantes pour l'analyse",
                'service_not_started': 'Service not started' if not isFr else "Service non commenc\u00e9",
                'duplicate_payment': 'Duplicate payment' if not isFr else 'Paiement en double',
                'other': 'Other' if not isFr else 'Autre'
            }.get(data.reason, data.reason)

            html_content = f"""
            <html><body style="font-family: Arial, sans-serif; padding: 20px;">
                <h2 style="color: #a553be;">New Refund Request - {refund_ref}</h2>
                <h3>Personal Information</h3>
                <p><strong>Name:</strong> {data.firstName} {data.lastName}</p>
                <p><strong>Username:</strong> {data.username}</p>
                <p><strong>Email:</strong> {data.email}</p>
                <p><strong>Phone:</strong> {data.phone}</p>
                <p><strong>Address:</strong> {data.address}, {data.postalCode} {data.city}, {data.country}</p>
                <h3>Order Information</h3>
                <p><strong>Order Reference:</strong> {data.orderReference}</p>
                <p><strong>Order Date:</strong> {data.orderDate}</p>
                <p><strong>Package:</strong> {data.packagePurchased}</p>
                <p><strong>Amount:</strong> &euro;{data.amountPaid}</p>
                <h3>Bank Information</h3>
                <p><strong>Account Holder:</strong> {data.accountHolder}</p>
                <p><strong>IBAN:</strong> {data.iban}</p>
                <p><strong>BIC:</strong> {data.bic}</p>
                <p><strong>Bank:</strong> {data.bankName}</p>
                <h3>Reason</h3>
                <p><strong>Reason:</strong> {reason_text}</p>
                <p><strong>Details:</strong> {data.additionalDetails or 'N/A'}</p>
                <hr><p style="color: #666;">Submitted: {data.submittedAt}</p>
            </body></html>
            """
            try:
                resend.Emails.send({
                    "from": "2good2breal <noreply@2good2breal.com>",
                    "to": admin_email,
                    "subject": f"[REFUND REQUEST] {refund_ref} - {data.firstName} {data.lastName}",
                    "html": html_content
                })
            except Exception as e:
                logger.error(f"Failed to send refund notification email: {e}")

        return {"message": "Refund request submitted successfully", "reference": refund_ref}
    except Exception as e:
        logger.error(f"Error submitting refund request: {e}")
        raise HTTPException(status_code=500, detail="Failed to submit refund request")


@router.post("/admin/analyses/{analysis_id}/report")
async def save_admin_report(analysis_id: str, data: AdminReportData, admin: dict = Depends(get_admin_user)):
    result = await db.verification_results.update_one(
        {"id": analysis_id},
        {"$set": {"admin_report": data.admin_report, "status": data.status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"message": "Report saved successfully"}


@router.delete("/admin/analyses/{analysis_id}")
async def admin_delete_analysis(analysis_id: str, admin: dict = Depends(get_admin_user)):
    analysis = await db.verification_results.find_one({"id": analysis_id})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    result = await db.verification_results.delete_one({"id": analysis_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=500, detail="Failed to delete analysis")
    logger.info(f"Admin deleted analysis {analysis_id}")
    return {"message": "Analysis deleted successfully"}


@router.get("/admin/analyses/{analysis_id}/submission-pdf")
async def download_submission_pdf(analysis_id: str, admin: dict = Depends(get_admin_user)):
    from reportlab.platypus import Image as RLImage
    from reportlab.lib.units import inch

    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    form_data = analysis.get("form_data", {})
    ai_analysis = analysis.get("ai_analysis", {})

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    elements = []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#7c3aed'), alignment=1, spaceAfter=20)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#7c3aed'), spaceBefore=15, spaceAfter=10, underline=True)
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#666666'), fontName='Helvetica-Bold')
    value_style = ParagraphStyle('Value', parent=styles['Normal'], fontSize=11, textColor=colors.black)
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#888888'), alignment=1)

    elements.append(Paragraph("2good2breal", title_style))
    elements.append(Paragraph("Profile Verification Service - Submission Form", ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=12, alignment=1, textColor=colors.HexColor('#666666'))))
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(f"Date: {analysis.get('created_at', 'N/A')[:10] if analysis.get('created_at') else 'N/A'}", ParagraphStyle('Date', parent=styles['Normal'], fontSize=10, alignment=2)))
    elements.append(Spacer(1, 15))

    elements.append(Paragraph("CLIENT INFORMATION", section_style))
    client_data = [
        ["Name:", analysis.get("user_name", "-"), "Email:", form_data.get("client_email", analysis.get("user_email", "-"))],
        ["Age:", form_data.get("client_age", "-"), "Location:", form_data.get("client_location", "-")],
    ]
    client_table = Table(client_data, colWidths=[70, 150, 70, 150])
    client_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')), ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#555555')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(client_table)
    elements.append(Spacer(1, 15))

    elements.append(Paragraph("PROFILE INFORMATION", section_style))
    profile_data_rows = [
        ["Profile Name:", form_data.get("profile_name", "-"), "Full Real Name:", form_data.get("full_real_name", "-")],
        ["Gender:", (form_data.get("gender", "-") or "-").capitalize(), "Height:", form_data.get("height", "-")],
        ["Nationality:", form_data.get("nationality", "-"), "Shared Language:", form_data.get("language_of_communication", "-")],
        ["Marital Status:", form_data.get("assumed_marital_status", "-"), "Hobbies/Interests:", form_data.get("hobbies_interests", "-")],
        ["University:", form_data.get("university_college", "-"), "Years Attendance / Graduation:", form_data.get("years_attendance", "-")],
    ]
    profile_table = Table(profile_data_rows, colWidths=[90, 130, 90, 130])
    profile_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')), ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#555555')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(profile_table)
    elements.append(Spacer(1, 15))

    elements.append(Paragraph("PROFILE DETAILS", section_style))
    details_data = [
        ["Date of Birth:", form_data.get("date_of_birth", "-"), "Known Age:", form_data.get("assumed_age", "-")],
        ["Location:", form_data.get("profile_location", "-"), "Platform:", form_data.get("dating_platform", "-")],
        ["Occupation:", form_data.get("occupation", "-"), "Company:", form_data.get("company_name", "-")],
    ]
    details_table = Table(details_data, colWidths=[90, 130, 90, 130])
    details_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')), ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#555555')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(details_table)
    elements.append(Spacer(1, 10))

    if form_data.get("profile_bio"):
        elements.append(Paragraph("Profile Bio:", label_style))
        elements.append(Paragraph(form_data.get("profile_bio", "-"), value_style))
        elements.append(Spacer(1, 15))

    elements.append(Paragraph("PHOTOS AND SOCIAL MEDIA", section_style))
    photos_data = [["Number of Photos:", str(form_data.get("profile_photos_count", len(form_data.get("photos", [])))), "Verified Photos:", "Yes" if form_data.get("has_verified_photos") else "No"]]
    photos_table = Table(photos_data, colWidths=[100, 120, 100, 120])
    photos_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10), ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(photos_table)

    if form_data.get("social_media_links"):
        elements.append(Paragraph("Social Media Links (User names):", label_style))
        elements.append(Paragraph(form_data.get("social_media_links", "-"), value_style))
    elements.append(Spacer(1, 15))

    elements.append(Paragraph("COMMUNICATION ANALYSIS", section_style))
    if form_data.get("communication_frequency"):
        elements.append(Paragraph("Communication Frequency:", label_style))
        elements.append(Paragraph(form_data.get("communication_frequency", "-"), value_style))
    if form_data.get("message_substance"):
        elements.append(Paragraph("Message Substance:", label_style))
        elements.append(Paragraph(form_data.get("message_substance", "-"), value_style))
    elements.append(Spacer(1, 15))

    elements.append(Paragraph("OBSERVATIONS & CONCERNS", section_style))
    elements.append(Paragraph(form_data.get("observations_concerns", "-"), value_style))
    elements.append(Spacer(1, 20))

    if ai_analysis:
        elements.append(Paragraph("AI ANALYSIS RESULTS", section_style))
        score = ai_analysis.get("overall_score", 0)
        trust_level = ai_analysis.get("trust_level", "unknown").replace("_", " ").upper()
        elements.append(Paragraph(f"Trust Score: {score}/100 - {trust_level}", ParagraphStyle('Score', parent=styles['Normal'], fontSize=14, textColor=colors.HexColor('#dc2626') if score < 40 else colors.HexColor('#22c55e') if score >= 70 else colors.HexColor('#eab308'))))
        if ai_analysis.get("analysis_summary"):
            elements.append(Spacer(1, 10))
            elements.append(Paragraph("AI Summary:", label_style))
            elements.append(Paragraph(ai_analysis.get("analysis_summary", "-"), value_style))
        red_flags = ai_analysis.get("red_flags", [])
        if red_flags:
            elements.append(Spacer(1, 10))
            elements.append(Paragraph(f"Red Flags Detected ({len(red_flags)}):", label_style))
            for flag in red_flags:
                elements.append(Paragraph(f"- [{flag.get('severity', 'low').upper()}] {flag.get('category', '')}: {flag.get('description', '')}", value_style))
        elements.append(Spacer(1, 20))

    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0'), spaceAfter=15))
    elements.append(Paragraph("2good2breal - Profile Verification Service", footer_style))
    elements.append(Paragraph("contact@2good2breal.com | +33 (0) 7 67 92 55 45 | www.2good2breal.com", footer_style))

    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()

    profile_name = form_data.get("profile_name", "submission").replace(" ", "_")
    filename = f"submission_{profile_name}_{analysis_id[:8]}.pdf"
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})


@router.get("/admin/analyses/{analysis_id}/submission-docx")
async def download_submission_docx(analysis_id: str, admin: dict = Depends(get_admin_user)):
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    form_data = analysis.get("form_data", {})
    ai_analysis = analysis.get("ai_analysis", {})
    photos = form_data.get("photos", [])

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # PAGE 1: HEADER + CLIENT INFO + PROFILE INFO
    title = doc.add_heading('2good2breal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(124, 58, 237)

    subtitle = doc.add_paragraph('Profile Verification Service - Submission Form')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    created_at = analysis.get('created_at', '')
    date_str = created_at[:10] if created_at else 'N/A'
    date_para.add_run(f"Date: {date_str}")
    doc.add_paragraph()

    doc.add_heading('CLIENT INFORMATION', level=1)
    client_table = doc.add_table(rows=2, cols=4)
    client_table.style = 'Table Grid'
    client_data_rows = [
        ("NAME", analysis.get("user_name", "-"), "EMAIL", form_data.get("client_email", analysis.get("user_email", "-"))),
        ("AGE", form_data.get("client_age", "-"), "LOCATION", form_data.get("client_location", "-")),
    ]
    for i, row_data in enumerate(client_data_rows):
        cells = client_table.rows[i].cells
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(row_data[1]) if row_data[1] else "-"
        cells[2].text = row_data[2]
        cells[2].paragraphs[0].runs[0].font.bold = True
        cells[3].text = str(row_data[3]) if row_data[3] else "-"
    doc.add_paragraph()

    doc.add_heading('PROFILE INFORMATION', level=1)
    profile_table = doc.add_table(rows=5, cols=4)
    profile_table.style = 'Table Grid'
    profile_rows = [
        ("PROFILE NAME", form_data.get("profile_name", "-"), "FULL REAL NAME", form_data.get("full_real_name", "-")),
        ("GENDER", (form_data.get("gender", "-") or "-").capitalize(), "HEIGHT", form_data.get("height", "-")),
        ("NATIONALITY", form_data.get("nationality", "-"), "SHARED LANGUAGE", form_data.get("language_of_communication", "-")),
        ("MARITAL STATUS", form_data.get("assumed_marital_status", "-"), "HOBBIES / INTERESTS", form_data.get("hobbies_interests", "-")),
        ("UNIVERSITY / COLLEGE", form_data.get("university_college", "-"), "YEAR/S OF ATTENDANCE / GRADUATION", form_data.get("years_attendance", "-")),
    ]
    for i, row_data in enumerate(profile_rows):
        cells = profile_table.rows[i].cells
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(row_data[1]) if row_data[1] else "-"
        cells[2].text = row_data[2]
        cells[2].paragraphs[0].runs[0].font.bold = True
        cells[3].text = str(row_data[3]) if row_data[3] else "-"

    # PAGE 2: PROFILE DETAILS + BIO
    doc.add_page_break()
    doc.add_heading('PROFILE DETAILS', level=1)
    details_tbl = doc.add_table(rows=4, cols=4)
    details_tbl.style = 'Table Grid'
    det_rows = [
        ("DATE OF BIRTH", form_data.get("date_of_birth", "-"), "KNOWN AGE", form_data.get("assumed_age", "-")),
        ("LOCATION", form_data.get("profile_location", "-"), "PLATFORM", form_data.get("dating_platform", "-")),
        ("OCCUPATION", form_data.get("occupation", "-"), "COMPANY NAME", form_data.get("company_name", "-")),
        ("COMPANY WEBSITE", form_data.get("company_website", "-"), "", ""),
    ]
    for i, row_data in enumerate(det_rows):
        cells = details_tbl.rows[i].cells
        cells[0].text = row_data[0]
        if row_data[0]: cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(row_data[1]) if row_data[1] else "-"
        cells[2].text = row_data[2]
        if row_data[2]: cells[2].paragraphs[0].runs[0].font.bold = True
        cells[3].text = str(row_data[3]) if row_data[3] else "-"
    doc.add_paragraph()
    doc.add_heading('PROFILE BIO', level=1)
    doc.add_paragraph(form_data.get("profile_bio", "-") or "-")

    # PAGE 3: PHOTOS AND SOCIAL MEDIA
    doc.add_page_break()
    doc.add_heading('PHOTOS AND SOCIAL MEDIA', level=1)
    photos_info_table = doc.add_table(rows=1, cols=4)
    photos_info_table.style = 'Table Grid'
    photos_count = form_data.get("profile_photos_count", len(photos))
    verified = "Yes" if form_data.get("has_verified_photos") else "No"
    cells = photos_info_table.rows[0].cells
    cells[0].text = "NUMBER OF PHOTOS"; cells[0].paragraphs[0].runs[0].font.bold = True
    cells[1].text = str(photos_count)
    cells[2].text = "VERIFIED PHOTOS"; cells[2].paragraphs[0].runs[0].font.bold = True
    cells[3].text = verified
    doc.add_paragraph()
    social_heading = doc.add_paragraph()
    social_heading.add_run("SOCIAL MEDIA LINKS (USER NAMES):").bold = True
    doc.add_paragraph(form_data.get("social_media_links", "-") or "-")
    doc.add_paragraph()

    if photos:
        photos_heading = doc.add_paragraph()
        photos_heading.add_run(f"UPLOADED PHOTOS ({len(photos)})").bold = True
        for idx, photo in enumerate(photos):
            photo_name = photo.get("name", photo.get("filename", f"Photo {idx + 1}"))
            photo_data = photo.get("base64", photo.get("data", ""))
            if photo_data and isinstance(photo_data, str):
                try:
                    if "," in photo_data: photo_data = photo_data.split(",")[1]
                    image_bytes = base64.b64decode(photo_data)
                    image_stream = BytesIO(image_bytes)
                    doc.add_picture(image_stream, width=Inches(3))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(photo_name)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.size = Pt(9); run.font.italic = True
                except Exception:
                    doc.add_paragraph(f"[Photo: {photo_name} - Could not embed]")
    doc.add_paragraph()

    doc.add_heading('ACTIVITY INFORMATION', level=1)
    activity_table = doc.add_table(rows=1, cols=4)
    activity_table.style = 'Table Grid'
    cells = activity_table.rows[0].cells
    cells[0].text = "PROFILE CREATION DATE"; cells[0].paragraphs[0].runs[0].font.bold = True
    cells[1].text = form_data.get("profile_creation_date", "-") or "-"
    cells[2].text = "LAST ACTIVE"; cells[2].paragraphs[0].runs[0].font.bold = True
    cells[3].text = form_data.get("last_active", "-") or "-"

    # PAGE 4: COMMUNICATION ANALYSIS + OBSERVATIONS
    doc.add_page_break()
    doc.add_heading('COMMUNICATION ANALYSIS', level=1)
    comm_heading = doc.add_paragraph()
    comm_heading.add_run("COMMUNICATION FREQUENCY:").bold = True
    doc.add_paragraph(form_data.get("communication_frequency", "-") or "-")
    msg_heading = doc.add_paragraph()
    msg_heading.add_run("MESSAGE SUBSTANCE:").bold = True
    doc.add_paragraph(form_data.get("message_substance", "-") or "-")
    doc.add_paragraph()
    doc.add_heading('OBSERVATIONS & CONCERNS', level=1)
    doc.add_paragraph(form_data.get("observations_concerns", "-") or "-")

    # PAGE 5: AI ANALYSIS
    if ai_analysis:
        doc.add_page_break()
        doc.add_heading('AI ANALYSIS RESULTS', level=1)
        score = ai_analysis.get("overall_score", 0)
        trust_level = ai_analysis.get("trust_level", "unknown").replace("_", " ").upper()
        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"Trust Score: {score}/100 - {trust_level}")
        score_run.font.size = Pt(16); score_run.font.bold = True
        if score < 40: score_run.font.color.rgb = RGBColor(220, 38, 38)
        elif score >= 70: score_run.font.color.rgb = RGBColor(34, 197, 94)
        else: score_run.font.color.rgb = RGBColor(234, 179, 8)
        doc.add_paragraph()
        if ai_analysis.get("analysis_summary"):
            summary_heading = doc.add_paragraph()
            summary_heading.add_run("AI SUMMARY:").bold = True
            doc.add_paragraph(ai_analysis.get("analysis_summary", "-"))
        doc.add_paragraph()
        red_flags = ai_analysis.get("red_flags", [])
        if red_flags:
            flags_heading = doc.add_paragraph()
            flags_heading.add_run(f"RED FLAGS DETECTED ({len(red_flags)}):").bold = True
            for flag in red_flags:
                doc.add_paragraph()
                flag_title = doc.add_paragraph()
                flag_title.add_run(f"{flag.get('category', 'Unknown')}:").bold = True
                desc_para = doc.add_paragraph()
                desc_para.add_run("Description: ").bold = True
                desc_para.add_run(flag.get('description', '-'))
                if flag.get('recommendation'):
                    rec_para = doc.add_paragraph()
                    rec_para.add_run("Recommendation: ").bold = True
                    rec_para.add_run(flag.get('recommendation', ''))
                sev_para = doc.add_paragraph()
                sev_para.add_run("Severity: ").bold = True
                sev_para.add_run(flag.get('severity', 'low').upper())
        else:
            doc.add_paragraph("No major red flags detected.")
        doc.add_paragraph()
        recommendations = ai_analysis.get("recommendations", [])
        if recommendations:
            rec_heading = doc.add_paragraph()
            rec_heading.add_run("AI RECOMMENDATIONS:").bold = True
            for rec in recommendations:
                doc.add_paragraph(f"- {rec}", style='List Bullet')

    # PAGE 6: RESEARCH AND VERIFICATIONS
    doc.add_page_break()
    doc.add_heading('Research and Verifications performed include some of the following:', level=1)
    verifications = [
        ("1. Platform Analysis", "Intense scrutinizing of all platforms used by 'the profile' in the past and present."),
        ("2. Occupation Verification", "Resourcing and authenticating profile's occupation via one on one discrete and direct communication means.\n\nAccess to occupation and / or company official website through various complex and often unattainable platforms.\n\nIntense cross-checking of the profile's email addresses and user names worldwide."),
        ("3. Photo Identification", "Photo identification via cross-checking of multiple image databases and reverse image search platforms.\n\nDetection and screening for multiple and stolen identities.\n\nResearch and confirmation of all Profile Platforms, Locations and Residencies previously and in use today."),
        ("4. Location Verification", "Verification of locations such as photo venues, background images and sceneries relating to where the profile claims to be or reside."),
        ("5. Location Cross Referencing", "Cross referencing of all the profile's locations and personal details to detect any mismatched information."),
        ("6. Photo Authenticity", "Clarity and authenticity of all photos provided by you and of those 2good2breal gains access to via websites, apps, platforms and other means."),
    ]
    for v_title, description in verifications:
        v_heading = doc.add_paragraph()
        v_heading.add_run(v_title).bold = True
        v_heading.runs[0].font.color.rgb = RGBColor(124, 58, 237)
        doc.add_paragraph(description)
        doc.add_paragraph()

    # PAGE 7: RECOMMENDATIONS + THANK YOU
    doc.add_page_break()
    doc.add_heading('Our Recommendations', level=1)
    doc.add_paragraph("Based on our analysis, we recommend:")
    our_recommendations = [
        "Block and report the account on the platform,",
        "Save evidence such as screenshots and user names in the event you need to report it in future,",
        "Talk to someone you trust about the situation for support if you feel the need.",
        "Consider stepping back or ending the conversation and /or contact.",
    ]
    for rec in our_recommendations:
        doc.add_paragraph(f"- {rec}", style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph("As the situation is extremely ambiguous, in our opinion, it is essential to walk away and disconnect.")
    doc.add_paragraph()
    additional_recs = [
        "If your situation with this profile has escalated to a point that you feel overwhelmed, do not hesitate to seek professional help.",
        "Keep your offline life grounded and intact.",
    ]
    for rec in additional_recs:
        doc.add_paragraph(f"- {rec}", style='List Bullet')
    doc.add_paragraph()
    further_para = doc.add_paragraph()
    further_para.add_run("If you wish further analyzing of this profile, please provide us with more personal details such as extended family information, presumed previous occupations and subsequent history on your next request.").italic = True
    doc.add_paragraph()
    doc.add_paragraph()
    thank_you = doc.add_heading('Thank you for choosing 2good2breal', level=1)
    thank_you.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in thank_you.runs:
        run.font.color.rgb = RGBColor(124, 58, 237)
    ty_text1 = doc.add_paragraph("We hope this report assists to clarify, confirm or dismiss any doubts you may have of your Profile's authenticity or intentions.")
    ty_text1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ty_text2 = doc.add_paragraph("Furthermore, our team aims to provide you with an objective, informative and reliable report to help guide you towards well founded and smart decision making with this person in future.")
    ty_text2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    all_best = doc.add_paragraph("All the best from our team at 2good2breal.")
    all_best.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in all_best.runs:
        run.font.bold = True; run.font.color.rgb = RGBColor(124, 58, 237)
    doc.add_paragraph()
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"Contact: contact@2good2breal.com\n")
    contact.add_run(f"Report Reference: {analysis_id}\n")
    contact.add_run("This analysis should not be considered as legal advice.").italic = True
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("2good2breal - Profile Verification Service\n").bold = True
    footer.add_run("contact@2good2breal.com | +33 (0) 7 67 92 55 45 | www.2good2breal.com\n")
    footer.add_run("This document is confidential.").italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    profile_name = form_data.get("profile_name", "submission").replace(" ", "_")
    filename = f"submission_{profile_name}_{analysis_id[:8]}.docx"
    return Response(content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={filename}"})


@router.post("/admin/analyses/{analysis_id}/send-report")
async def send_report_to_client(analysis_id: str, data: SendReportData, admin: dict = Depends(get_admin_user)):
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    user = await db.users.find_one({"id": analysis["user_id"]}, {"_id": 0, "email": 1, "name": 1})
    form_data = analysis.get("form_data", {})
    report = data.admin_report

    verdict_color = "#10b981" if report.get("verdict") == "safe" else "#f59e0b" if report.get("verdict") == "suspicious" else "#ef4444" if report.get("verdict") == "dangerous" else "#6b7280"
    verdict_text = "SAFE - Profile appears authentic" if report.get("verdict") == "safe" else "SUSPICIOUS - Exercise caution" if report.get("verdict") == "suspicious" else "DANGEROUS - High risk of scam" if report.get("verdict") == "dangerous" else "INCONCLUSIVE - More information needed"

    html_content = f"""
    <!DOCTYPE html>
    <html><head><meta charset="UTF-8">
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }}
        .header {{ border-bottom: 3px solid #0891b2; padding-bottom: 20px; margin-bottom: 30px; }}
        .header h1 {{ color: #0891b2; margin: 0; }}
        .section {{ margin-bottom: 30px; padding: 20px; border-radius: 8px; }}
        .section-title {{ font-size: 18px; font-weight: bold; color: #334155; margin-bottom: 15px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
        .info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }}
        .info-item {{ margin: 5px 0; }}
        .verdict-box {{ padding: 15px; border-radius: 8px; text-align: center; margin: 20px 0; }}
        .verdict-text {{ font-size: 24px; font-weight: bold; color: white; }}
        .analysis-box {{ background: #f8fafc; padding: 15px; border-radius: 8px; margin: 10px 0; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; text-align: center; color: #666; font-size: 12px; }}
    </style>
    </head><body>
        <div class="header">
            <h1>2good2breal</h1>
            <p style="color: #666; margin: 5px 0;">Profile Verification Report</p>
            <p style="font-size: 14px; color: #666;">Report Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')} | Reference: #{analysis_id[:8].upper()}</p>
        </div>
        <div class="section" style="background: #f0f9ff; border: 2px solid #0891b2;">
            <div class="section-title" style="color: #0891b2;">Expert Verification Verdict</div>
            <div class="verdict-box" style="background: {verdict_color};"><div class="verdict-text">{verdict_text}</div></div>
        </div>
        <div class="section" style="background: #f8fafc;">
            <div class="section-title">Profile Analyzed</div>
            <div class="info-grid">
                <div class="info-item"><strong>Profile Name:</strong> {form_data.get('profile_name', 'N/A')}</div>
                <div class="info-item"><strong>Platform:</strong> {form_data.get('dating_platform', 'N/A')}</div>
                <div class="info-item"><strong>Location:</strong> {form_data.get('profile_location', 'N/A')}</div>
                <div class="info-item"><strong>Occupation:</strong> {form_data.get('occupation', 'N/A')}</div>
            </div>
        </div>
        <div class="section"><div class="section-title">Detailed Analysis</div><div class="analysis-box">{report.get('detailedAnalysis', 'No detailed analysis provided.').replace(chr(10), '<br>')}</div></div>
        <div class="section"><div class="section-title">Our Recommendations</div><div class="analysis-box">{report.get('recommendations', 'No specific recommendations.').replace(chr(10), '<br>')}</div></div>
        {f'<div class="section"><div class="section-title">Additional Notes</div><div class="analysis-box">{report.get("additionalNotes", "").replace(chr(10), "<br>")}</div></div>' if report.get('additionalNotes') else ''}
        <div class="footer">
            <p>This report was generated by 2good2breal verification service.</p>
            <p>For questions, contact: contact@2good2breal.com</p>
            <p style="font-style: italic; margin-top: 15px;">This analysis is based on information provided and should not be considered as legal advice.</p>
        </div>
    </body></html>
    """

    try:
        resend.api_key = os.environ.get("RESEND_API_KEY")
        resend.emails.send({
            "from": "2good2breal <noreply@resend.dev>",
            "to": [data.client_email],
            "subject": f"Your Profile Verification Report - {form_data.get('profile_name', 'Analysis')}",
            "html": html_content
        })
        await db.verification_results.update_one(
            {"id": analysis_id},
            {"$set": {"admin_report": data.admin_report, "status": "completed", "report_sent_at": datetime.now(timezone.utc).isoformat(), "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        logger.info(f"Report sent to {data.client_email} for analysis {analysis_id}")
        return {"message": "Report sent successfully"}
    except Exception as e:
        logger.error(f"Failed to send report: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")


@router.get("/admin/analyses/{analysis_id}/download-docx")
async def download_report_docx(analysis_id: str, admin: dict = Depends(get_admin_user)):
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    user = await db.users.find_one({"id": analysis.get("user_id")}, {"_id": 0, "email": 1, "name": 1})
    if user:
        analysis["user_email"] = user.get("email", "")
        analysis["user_name"] = user.get("name", "")

    admin_report = analysis.get("admin_report", {})
    docx_bytes = generate_report_docx(analysis, admin_report)

    profile_name = analysis.get("form_data", {}).get("profile_name", "report")
    safe_name = "".join(c if c.isalnum() or c in "._- " else "_" for c in profile_name)
    filename = f"Report_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"

    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
