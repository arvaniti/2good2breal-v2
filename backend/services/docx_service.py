import logging
import base64
from io import BytesIO
from datetime import datetime, timezone
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from config import ROOT_DIR

logger = logging.getLogger(__name__)


def generate_report_docx(analysis: dict, admin_report: dict) -> bytes:
    """Generate a DOCX report matching the EXACT template layout."""
    doc = Document()

    form_data = analysis.get("form_data", {})
    ai = analysis.get("ai_analysis", {})

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    logo_path = ROOT_DIR / "logo.png"
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    if logo_path.exists():
        try:
            run = header_para.add_run()
            run.add_picture(str(logo_path), width=Inches(1.2))
        except Exception as e:
            logger.warning(f"Logo error: {e}")

    def style_header_cell(cell, text):
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)

    def style_data_cell(cell, text):
        cell.text = str(text) if text else "-"

    # PAGE 1
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("2good2breal")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(124, 58, 237)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Profile Verification Report")
    sub_run.bold = True
    sub_run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    # CLIENT INFORMATION
    client_header = doc.add_paragraph()
    client_run = client_header.add_run("CLIENT INFORMATION")
    client_run.bold = True
    client_run.font.size = Pt(12)

    client_table = doc.add_table(rows=2, cols=4)
    client_table.style = 'Table Grid'
    style_header_cell(client_table.rows[0].cells[0], "NAME")
    style_data_cell(client_table.rows[0].cells[1], analysis.get("user_name", ""))
    style_header_cell(client_table.rows[0].cells[2], "EMAIL")
    style_data_cell(client_table.rows[0].cells[3], analysis.get("user_email", ""))
    style_header_cell(client_table.rows[1].cells[0], "AGE")
    style_data_cell(client_table.rows[1].cells[1], form_data.get("client_age", ""))
    style_header_cell(client_table.rows[1].cells[2], "LOCATION")
    style_data_cell(client_table.rows[1].cells[3], form_data.get("client_location", ""))
    doc.add_paragraph()

    # PROFILE INFORMATION
    profile_header = doc.add_paragraph()
    profile_run = profile_header.add_run("PROFILE INFORMATION")
    profile_run.bold = True
    profile_run.font.size = Pt(12)

    profile_table = doc.add_table(rows=5, cols=4)
    profile_table.style = 'Table Grid'
    style_header_cell(profile_table.rows[0].cells[0], "PROFILE NAME")
    style_data_cell(profile_table.rows[0].cells[1], form_data.get("profile_name", ""))
    style_header_cell(profile_table.rows[0].cells[2], "FULL REAL NAME")
    style_data_cell(profile_table.rows[0].cells[3], form_data.get("full_real_name", ""))
    style_header_cell(profile_table.rows[1].cells[0], "GENDER")
    gender = form_data.get("gender", "") or ""
    style_data_cell(profile_table.rows[1].cells[1], gender.capitalize() if gender else "")
    style_header_cell(profile_table.rows[1].cells[2], "HEIGHT")
    style_data_cell(profile_table.rows[1].cells[3], form_data.get("height", ""))
    style_header_cell(profile_table.rows[2].cells[0], "NATIONALITY")
    style_data_cell(profile_table.rows[2].cells[1], form_data.get("nationality", ""))
    style_header_cell(profile_table.rows[2].cells[2], "SHARED LANGUAGE")
    shared_lang = form_data.get("language_of_communication", "") or form_data.get("shared_language", "") or ""
    style_data_cell(profile_table.rows[2].cells[3], shared_lang)
    style_header_cell(profile_table.rows[3].cells[0], "MARITAL STATUS")
    style_data_cell(profile_table.rows[3].cells[1], form_data.get("assumed_marital_status", ""))
    style_header_cell(profile_table.rows[3].cells[2], "HOBBIES / INTERESTS")
    style_data_cell(profile_table.rows[3].cells[3], form_data.get("hobbies_interests", ""))
    style_header_cell(profile_table.rows[4].cells[0], "UNIVERSITY / COLLEGE")
    style_data_cell(profile_table.rows[4].cells[1], form_data.get("university_college", ""))
    style_header_cell(profile_table.rows[4].cells[2], "YEAR/S OF ATTENDANCE / GRADUATION")
    style_data_cell(profile_table.rows[4].cells[3], form_data.get("years_of_attendance", ""))
    doc.add_paragraph()

    # PROFILE DETAILS
    details_header = doc.add_paragraph()
    details_run = details_header.add_run("PROFILE DETAILS")
    details_run.bold = True
    details_run.font.size = Pt(12)

    details_table = doc.add_table(rows=4, cols=4)
    details_table.style = 'Table Grid'
    style_header_cell(details_table.rows[0].cells[0], "DATE OF BIRTH")
    style_data_cell(details_table.rows[0].cells[1], form_data.get("date_of_birth", ""))
    style_header_cell(details_table.rows[0].cells[2], "KNOWN AGE")
    style_data_cell(details_table.rows[0].cells[3], form_data.get("assumed_age", ""))
    style_header_cell(details_table.rows[1].cells[0], "LOCATION")
    style_data_cell(details_table.rows[1].cells[1], form_data.get("profile_location", ""))
    style_header_cell(details_table.rows[1].cells[2], "PLATFORM")
    style_data_cell(details_table.rows[1].cells[3], form_data.get("dating_platform", ""))
    style_header_cell(details_table.rows[2].cells[0], "OCCUPATION")
    style_data_cell(details_table.rows[2].cells[1], form_data.get("occupation", ""))
    style_header_cell(details_table.rows[2].cells[2], "COMPANY NAME")
    style_data_cell(details_table.rows[2].cells[3], form_data.get("company_name", ""))
    style_header_cell(details_table.rows[3].cells[0], "COMPANY WEBSITE")
    style_data_cell(details_table.rows[3].cells[1], form_data.get("company_website", ""))
    details_table.rows[3].cells[2].text = ""
    details_table.rows[3].cells[3].text = ""

    # PAGE 2: ANALYSIS RESULTS
    doc.add_page_break()
    results_header = doc.add_paragraph()
    results_run = results_header.add_run("ANALYSIS RESULTS")
    results_run.bold = True
    results_run.font.size = Pt(14)

    score = ai.get("overall_score", 0) if ai else 0
    if score <= 25: risk_level = "EXTREME HIGH RISK"
    elif score <= 51: risk_level = "HIGH"
    elif score <= 70: risk_level = "MEDIUM"
    elif score <= 85: risk_level = "LOW"
    else: risk_level = "VERY LOW"

    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_para.add_run(f"Trust Score: {score}/100 - {risk_level}")
    score_run.bold = True
    score_run.font.size = Pt(16)
    doc.add_paragraph()

    rating_path = ROOT_DIR / "rating_scale.png"
    if rating_path.exists():
        try:
            doc.add_picture(str(rating_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except:
            pass

    summary = ai.get("analysis_summary", "") or ai.get("summary", "") if ai else ""
    if summary:
        s_header = doc.add_paragraph()
        s_header.add_run("SUMMARY").bold = True
        doc.add_paragraph(summary)
        doc.add_paragraph()

    red_flags = ai.get("red_flags", []) if ai else []
    rf_header = doc.add_paragraph()
    rf_header.add_run(f"RED FLAGS DETECTED ({len(red_flags)})").bold = True

    for flag in red_flags:
        doc.add_paragraph()
        if isinstance(flag, dict):
            cat = flag.get("category", "") or flag.get("type", "Unknown")
            cat_para = doc.add_paragraph()
            cat_para.add_run(cat).bold = True
            if flag.get("description"):
                d = doc.add_paragraph()
                d.add_run("Description: ").bold = True
                d.add_run(flag["description"])
            if flag.get("recommendation"):
                r = doc.add_paragraph()
                r.add_run("Recommendation: ").bold = True
                r.add_run(flag["recommendation"])
            sev = (flag.get("severity", "LOW") or "LOW").upper()
            sv = doc.add_paragraph()
            sv.add_run("Severity: ").bold = True
            sv.add_run(sev)
        else:
            doc.add_paragraph(f"- {str(flag)}")

    doc.add_paragraph()
    reco_h = doc.add_paragraph()
    reco_h.add_run("SOME RECOMMENDATIONS").bold = True
    ai_recs = ai.get("recommendations", []) if ai else []
    recs = ai_recs if ai_recs else [
        "Continue communicating through the platform or verified channels.",
        "Schedule a video call to fully bridge the gap between digital profile and reality.",
        "The lack of news for one day is common; do not interpret this as a 'disappearing' tactic yet.",
        "Verify 'travel' claims if financial assistance is requested."
    ]
    for rec in recs:
        doc.add_paragraph(f"- {rec}")

    # PAGE 3: CONCLUSIVE ANALYSIS - POINTS
    doc.add_page_break()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("CONCLUSIVE ANALYSIS - POINTS").bold = True

    # PAGE 4: CONCLUSIVE ANALYSIS - OVERALL
    doc.add_page_break()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("CONCLUSIVE ANALYSIS - OVERALL").bold = True

    # PAGE 5: RECOMMANDATIONS OVERALL
    doc.add_page_break()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run("RECOMMANDATIONS OVERALL").bold = True

    # PAGE 6: Research and Verifications
    doc.add_page_break()
    research_h = doc.add_paragraph()
    research_h.add_run("Research and Verifications performed include some of the following:").bold = True
    doc.add_paragraph()

    verifications = [
        ("1. Platform Analysis", "Intense scrutinizing of all platforms used by 'the profile' in the past and present."),
        ("2. Occupation Verification", "Resourcing and authenticating profile's occupation via direct communication means."),
        ("3. Photo Identification", "Photo identification via cross-checking of multiple image databases and reverse image search platforms.\n\nResearch and confirmation of all Profile Platforms, Locations and Residencies previously and in use today."),
        ("4. Location Verification", "Verification of locations such as photo venues, background images and sceneries."),
        ("5. Location Cross Referencing", "Cross referencing of all the profile's locations and personal details."),
        ("6. Photo Authenticity", "Clarity and authenticity of all photos provided by you and accessed via various means.")
    ]
    for title, desc in verifications:
        vp = doc.add_paragraph()
        vp.add_run(f"{title} ").bold = True
        vp.add_run(desc)
        doc.add_paragraph()

    # PAGE 7: Additional Recommendations
    doc.add_page_break()
    add_h = doc.add_paragraph()
    add_h.add_run("Additional Recommendations").bold = True
    doc.add_paragraph()

    add_recs = [
        "Block and report the account on the platform,",
        "Save evidence such as screenshots and user names for future reference,",
        "Talk to someone you trust about the situation for support,",
        "Consider stepping back or ending the conversation and/or contact,",
        "If overwhelmed, do not hesitate to seek professional help,",
        "Keep your offline life grounded and intact."
    ]
    for ar in add_recs:
        doc.add_paragraph(f"- {ar}")

    doc.add_paragraph()
    doc.add_paragraph("If you wish further analyzing of this profile, please provide us with more personal details such as extended family information, presumed previous occupations and subsequent history on your next request.")
    doc.add_paragraph()

    thank_p = doc.add_paragraph()
    thank_p.add_run("Thank you for choosing 2good2breal").bold = True
    doc.add_paragraph()
    doc.add_paragraph("We hope this report assists to clarify, confirm or dismiss any doubts you may have of your Profile's authenticity or intentions. All the best from our team at 2good2breal.")
    doc.add_paragraph()

    contact_p = doc.add_paragraph()
    contact_p.add_run("Contact: ").bold = True
    contact_p.add_run("contact@2good2breal.com")
    ref_p = doc.add_paragraph()
    ref_p.add_run("Report Reference: ").bold = True
    ref_p.add_run(analysis.get("id", ""))
    doc.add_paragraph()

    disc_p = doc.add_paragraph()
    disc_p.add_run("This analysis should not be considered as legal advice.").bold = True
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("2good2breal - Profile Verification Service\n").bold = True
    footer_p.add_run("contact@2good2breal.com | +33 (0) 7 67 92 55 45 | www.2good2breal.com")
    doc.add_paragraph()

    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_p.add_run("This document is confidential.").bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
